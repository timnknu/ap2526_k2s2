import docx

new_doc = docx.Document()

n_rows_we_want = 10

new_doc.add_paragraph().add_run('Here is my new small table')


d = docx.Document('big_table.docx')
tbl = d.tables[0]
for J, row in enumerate(tbl.rows):
    print(row.cells[0].text)
    if J % n_rows_we_want == 0:
        new_doc.add_paragraph().add_run('New part')
        t_new = new_doc.add_table(rows=0, cols=len(tbl.columns))
        if J>0:
            new_row = t_new.add_row()
            for i, src_cell in enumerate(tbl.rows[0].cells):
                new_row.cells[i].text = src_cell.text

    #
    new_row = t_new.add_row()
    for i, src_cell in enumerate(row.cells):
        new_row.cells[i].text = src_cell.text
    #


new_doc.save('generated_table.docx')
