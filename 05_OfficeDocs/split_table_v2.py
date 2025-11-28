import docx

d = docx.Document('big_table.docx')
body = d._element[0]
print(body)
print(body.getchildren())
table = body[1]
print(table.getchildren())

new_doc = docx.Document()
new_doc.add_paragraph().add_run('Here is my new small table')
t = new_doc.add_table(rows=0, cols=0)
t._element[0:5] = table[0:5]
new_doc.save('new_table_2.docx')